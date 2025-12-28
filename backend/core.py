import os
import asyncio
import re
import logging
import tempfile
import shutil
from datetime import datetime
from urllib.parse import urlparse

import google.generativeai as genai
from dotenv import load_dotenv
from fpdf import FPDF
from docx import Document

# Setup Logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class TranscribeCore:
    def __init__(self, api_key: str = None):
        # Initialize Gemini
        if not api_key:
            # Force reload .env to pick up changes
            load_dotenv(override=True)
            api_key = os.getenv("GEMINI_API_KEY")

        if api_key:
            genai.configure(api_key=api_key)
            self.model = genai.GenerativeModel('gemini-2.0-flash')
            masked_key = api_key[:4] + "..." + api_key[-4:] if len(api_key) > 8 else "***"
            logger.info(f"Gemini API initialized successfully with key: {masked_key}")
        else:
            logger.warning("GEMINI_API_KEY not found in environment variables.")
            self.model = None

    def _extract_shortcode(self, url: str) -> str:
        """Extract shortcode from Instagram URL."""
        parsed = urlparse(url)
        path = parsed.path.strip('/')
        parts = path.split('/')
        
        for i, part in enumerate(parts):
            if part in ['reel', 'p', 'reels'] and i + 1 < len(parts):
                return parts[i + 1]
        
        return parts[-1] if parts else "unknown"

    async def process_link(self, link: str, progress_callback=None):
        """Main processing logic for a single link."""
        temp_dir = None
        try:
            shortcode = self._extract_shortcode(link)
            
            if progress_callback:
                await progress_callback(f"Processing reel: {shortcode}", "info")

            # 1. Download the Instagram reel using yt-dlp
            if progress_callback:
                await progress_callback("Downloading reel using yt-dlp...", "info")
            
            temp_dir = tempfile.mkdtemp(prefix="transcribe_")
            video_path = await self._download_with_ytdlp(link, temp_dir, progress_callback)
            
            if not video_path:
                raise Exception("Failed to download video. The reel may be private or unavailable.")

            # 2. Generate Transcript via Gemini
            if progress_callback:
                await progress_callback("Sending video to Gemini AI for transcription...", "info")
            
            transcript_text = await self._transcribe_with_gemini(video_path, progress_callback)
            
            if not transcript_text:
                transcript_text = "Transcription failed or returned empty."

            # 3. Generate Files
            if progress_callback:
                await progress_callback("Generating PDF and DOCX files...", "info")
            
            files = self._generate_files(shortcode, link, transcript_text)
            
            if progress_callback:
                await progress_callback("Files generated successfully!", "success")

            return {
                "shortcode": shortcode,
                "preview": transcript_text[:200] + "..." if len(transcript_text) > 200 else transcript_text,
                "files": files,
                "full_transcript": transcript_text
            }

        except Exception as e:
            logger.error(f"Error processing link {link}: {e}")
            if progress_callback:
                await progress_callback(f"Error: {str(e)}", "error")
            raise e
        finally:
            if temp_dir and os.path.exists(temp_dir):
                try:
                    shutil.rmtree(temp_dir)
                except:
                    pass

    async def _download_with_ytdlp(self, url: str, temp_dir: str, progress_callback=None) -> str:
        """Download Instagram reel using yt-dlp."""
        try:
            import yt_dlp
            
            output_template = os.path.join(temp_dir, '%(id)s.%(ext)s')
            
            ydl_opts = {
                'outtmpl': output_template,
                'format': 'best[ext=mp4]/best',
                'quiet': True,
                'no_warnings': True,
                'extract_flat': False,
                'user_agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            }
            
            loop = asyncio.get_event_loop()
            
            def download_sync():
                try:
                    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                        ydl.download([url])
                    return True
                except Exception as e:
                    logger.error(f"yt-dlp error: {e}")
                    return False
            
            if progress_callback:
                await progress_callback("yt-dlp downloading video...", "info")
            
            success = await loop.run_in_executor(None, download_sync)
            
            if not success:
                return None
            
            # Find downloaded video
            for file in os.listdir(temp_dir):
                if file.endswith(('.mp4', '.webm', '.mkv')):
                    video_path = os.path.join(temp_dir, file)
                    file_size = os.path.getsize(video_path) / (1024 * 1024)  # MB
                    if progress_callback:
                        await progress_callback(f"Video downloaded: {file} ({file_size:.1f} MB)", "info")
                    return video_path
            
            return None
            
        except ImportError:
            if progress_callback:
                await progress_callback("yt-dlp not installed. Run: pip install yt-dlp", "error")
            return None
        except Exception as e:
            logger.error(f"Download error: {e}")
            if progress_callback:
                await progress_callback(f"Download failed: {str(e)}", "error")
            return None

    async def _transcribe_with_gemini(self, video_path: str, progress_callback=None) -> str:
        """Transcribe video using brute-force model selection."""
        api_key = os.getenv("GEMINI_API_KEY") 
        if not api_key:
             return "Error: Gemini API Key not configured."

        # Verified models from check_models.py
        models_to_try = [
            "gemini-2.5-flash",
            "gemini-2.5-pro",
            "gemini-2.0-flash",
            "gemini-2.0-flash-exp",
            "gemini-flash-latest",
        ]

        uploaded_file = None
        
        try:
            if progress_callback:
                await progress_callback("Uploading video to Gemini...", "info")
            
            # Upload video once
            uploaded_file = genai.upload_file(video_path)
            
            # Wait for processing
            while uploaded_file.state.name == "PROCESSING":
                await asyncio.sleep(1)
                uploaded_file = genai.get_file(uploaded_file.name)
            
            if uploaded_file.state.name == "FAILED":
                raise Exception("Video processing failed on Gemini servers during upload.")

            # Try models
            prompt = """Transcribe the audio from this video with high accuracy.
Instructions:
1. Detect and transcribe ANY language (English, Hindi, Tamil, etc.).
2. Handle code-switching (e.g., "Hinglish") exactly as spoken.
3. Format: 'Speaker: [Text]'.
4. If no speech, describe audio context."""

            last_error = None

            for model_name in models_to_try:
                try:
                    if progress_callback:
                        await progress_callback(f"Taking a shot with {model_name}...", "info")
                    
                    model = genai.GenerativeModel(model_name)
                    response = model.generate_content([uploaded_file, prompt])
                    
                    if response and response.text:
                        if progress_callback:
                            await progress_callback(f"Success with {model_name}!", "success")
                        return response.text
                
                except Exception as e:
                    error_msg = str(e)
                    logger.warning(f"{model_name} failed: {error_msg}")
                    if progress_callback:
                        await progress_callback(f"{model_name} failed: {error_msg}", "warning")
                    last_error = error_msg
                    # If it's a legitimate rate limit (429), we might want to wait, 
                    # but for "brute force" we just move to the next model which might have quota.
                    continue

            return f"All models failed. Last error: {last_error}"

        except Exception as e:
            logger.error(f"Fatal transcription error: {e}")
            return f"Fatal error: {str(e)}"
        
        finally:
            # Cleanup
            if uploaded_file:
                try:
                    genai.delete_file(uploaded_file.name)
                except:
                    pass

    def _generate_files(self, shortcode: str, link: str, text: str):
        """Generates PDF and DOCX files."""
        base_name = re.sub(r'[<>:"/\\|?*]', '', shortcode)
        if not base_name:
            base_name = "transcript"
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = "outputs"
        os.makedirs(output_dir, exist_ok=True)

        # PDF
        pdf_filename = f"{base_name}_{timestamp}.pdf"
        pdf_path = os.path.join(output_dir, pdf_filename)
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", 'B', size=16)
        pdf.cell(0, 10, "Instagram Reel Transcription", ln=True, align='C')
        pdf.set_font("Arial", size=10)
        pdf.cell(0, 10, f"Source: {link[:80]}...", ln=True)
        pdf.cell(0, 10, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
        pdf.ln(10)
        pdf.set_font("Arial", size=12)
        try:
            pdf.multi_cell(0, 10, text.encode('latin-1', 'replace').decode('latin-1'))
        except:
            pdf.multi_cell(0, 10, text.encode('ascii', 'replace').decode('ascii'))
        pdf.output(pdf_path)

        # DOCX
        docx_filename = f"{base_name}_{timestamp}.docx"
        docx_path = os.path.join(output_dir, docx_filename)
        doc = Document()
        doc.add_heading('Instagram Reel Transcription', 0)
        doc.add_paragraph(f"Source: {link}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")
        doc.add_heading('Transcription', level=1)
        doc.add_paragraph(text)
        doc.save(docx_path)

        base_url = "http://10.0.2.2:8000/files"
        return {
            "pdf": f"{base_url}/{pdf_filename}",
            "docx": f"{base_url}/{docx_filename}"
        }
